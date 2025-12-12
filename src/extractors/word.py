"""
Word document extractor - Extract text and images from .docx files
"""

import logging
from pathlib import Path
from typing import Optional

from extractors.base import BaseExtractor, ExtractionResult, ExtractionInterrupted

logger = logging.getLogger(__name__)


class WordExtractor(BaseExtractor):
    """Extract text and images from Word documents (.docx)"""
    
    SUPPORTED_EXTENSIONS = ['.docx']
    
    def __init__(self, output_base_path: Path):
        super().__init__(output_base_path)
        self.docx_available = self._check_docx()
    
    def _check_docx(self) -> bool:
        """Check if python-docx is available"""
        try:
            import docx
            return True
        except ImportError:
            logger.warning("python-docx not available - Word document extraction disabled")
            return False
    
    def can_extract(self, filepath: Path) -> bool:
        """Check if this extractor can handle the given file"""
        return filepath.suffix.lower() == '.docx' and self.docx_available
    
    def extract(self, filepath: Path, output_dir: Path) -> ExtractionResult:
        """
        Extract text and images from Word document
        
        Args:
            filepath: Path to Word document
            output_dir: Directory for extracted files
            
        Returns:
            ExtractionResult with extraction details
        """
        result = ExtractionResult(filepath)
        
        if not self.docx_available:
            result.add_error("python-docx not installed - cannot process Word documents")
            return result
        
        try:
            import docx
            from docx.oxml import parse_xml
            from PIL import Image
            import io
            
            logger.info(f"Extracting Word document: {filepath.name}")
            
            # Ensure output directory exists
            if not self.ensure_output_dir(output_dir):
                result.add_error(f"Failed to create output directory: {output_dir}")
                return result
            
            # Create subdirectory for images
            file_safe_name = self.sanitize_filename(filepath.name)
            images_dir = output_dir / f"{file_safe_name}_images"
            
            # Open document
            doc = docx.Document(filepath)
            
            # Report substep
            self.report_substep("Extracting text content")
            
            # Extract text
            text_output = output_dir / f"{file_safe_name}.txt"
            text_content = self._extract_text(doc, result)
            
            if text_content.strip():
                with open(text_output, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                result.add_file(text_output)
                logger.info(f"Extracted text to {text_output.name}")
            else:
                result.add_warning("No text content found in document")
            
            # Check for interrupt before image extraction
            self.check_interrupted()
            
            # Extract images
            self.report_substep("Extracting embedded images")
            image_count = self._extract_images(doc, images_dir, result)
            
            if image_count > 0:
                result.metadata['images_extracted'] = image_count
                logger.info(f"Extracted {image_count} images")
            else:
                logger.info("No images found in document")
            
            if len(result.extracted_files) > 0:
                result.success = True
                logger.info(f"Successfully extracted data from {filepath.name}")
            else:
                result.add_warning("No data extracted from Word document")
        
        except ExtractionInterrupted:
            # Re-raise interrupt exceptions so they propagate to the manager
            raise
            
        except Exception as e:
            error_msg = f"Failed to extract {filepath.name}: {str(e)}"
            logger.error(error_msg)
            result.add_error(error_msg)
        
        return result
    
    def _extract_text(self, doc, result: ExtractionResult) -> str:
        """Extract all text from Word document"""
        try:
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        text_parts.append(f"\n{'='*80}\n")
                        text_parts.append(f"{text}\n")
                        text_parts.append(f"{'='*80}\n\n")
                    else:
                        text_parts.append(f"{text}\n\n")
            
            # Extract text from tables
            if doc.tables:
                text_parts.append(f"\n{'='*80}\n")
                text_parts.append("TABLES\n")
                text_parts.append(f"{'='*80}\n\n")
                
                for table_idx, table in enumerate(doc.tables, 1):
                    text_parts.append(f"--- Table {table_idx} ---\n")
                    
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        if row_text:
                            text_parts.append(f"{row_text}\n")
                    
                    text_parts.append("\n")
            
            result.metadata['paragraph_count'] = len(doc.paragraphs)
            result.metadata['table_count'] = len(doc.tables)
            
            return ''.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            result.add_warning(f"Text extraction error: {e}")
            return ""
    
    def _extract_images(self, doc, output_dir: Path, result: ExtractionResult) -> int:
        """Extract all images from Word document"""
        try:
            from PIL import Image
            import io
            
            image_count = 0
            
            # Get all image relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Ensure output directory exists
                        if not self.ensure_output_dir(output_dir):
                            result.add_warning(f"Failed to create images directory: {output_dir}")
                            continue
                        
                        # Get image data
                        image_data = rel.target_part.blob
                        
                        # Determine image extension
                        content_type = rel.target_part.content_type
                        if 'png' in content_type:
                            ext = 'png'
                        elif 'jpeg' in content_type or 'jpg' in content_type:
                            ext = 'jpg'
                        elif 'gif' in content_type:
                            ext = 'gif'
                        elif 'bmp' in content_type:
                            ext = 'bmp'
                        else:
                            ext = 'png'  # default
                        
                        # Create filename
                        image_count += 1
                        img_filename = f"image{image_count}.{ext}"
                        img_path = output_dir / img_filename
                        
                        # Save image
                        with open(img_path, 'wb') as f:
                            f.write(image_data)
                        
                        result.add_file(img_path)
                        logger.debug(f"Extracted image: {img_filename}")
                        
                    except Exception as e:
                        logger.warning(f"Failed to extract image {image_count}: {e}")
            
            return image_count
            
        except Exception as e:
            logger.error(f"Error extracting images: {e}")
            result.add_warning(f"Image extraction error: {e}")
            return 0

